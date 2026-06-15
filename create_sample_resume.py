"""Generate a sample DOCX resume for testing."""

from docx import Document
import os

doc = Document()
doc.add_heading("John Smith", 0)
doc.add_paragraph("Email: john.smith@email.com | Phone: +1 (555) 123-4567")

doc.add_heading("Skills", level=1)
doc.add_paragraph("Python, JavaScript, React, Flask, SQL, Docker, Git, AWS, Machine Learning")

doc.add_heading("Education", level=1)
doc.add_paragraph("B.S. Computer Science, State University, 2020")
doc.add_paragraph("M.S. Data Science, Tech Institute, 2022")

doc.add_heading("Experience", level=1)
doc.add_paragraph("Senior Software Engineer | Tech Corp | 2022 - Present")
doc.add_paragraph("- Built REST APIs with Flask and Python serving 10K+ daily users")
doc.add_paragraph("- Led migration to Docker and AWS, reducing deployment time by 60%")

doc.add_paragraph("Software Developer | Startup Inc | 2020 - 2022")
doc.add_paragraph("- Developed React frontend components for customer dashboard")

doc.add_heading("Projects", level=1)
doc.add_paragraph("Resume Parser App")
doc.add_paragraph("- Full-stack web app using Flask and SQLite for automated resume parsing")
doc.add_paragraph("E-Commerce Platform")
doc.add_paragraph("- Built with React and Node.js with payment integration")

output = os.path.join(os.path.dirname(__file__), "sample_resume.docx")
doc.save(output)
print(f"Sample resume saved to: {output}")
