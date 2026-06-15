"""
DOCX text extraction using python-docx.
Extracts raw text from uploaded Word resume files.
"""

from docx import Document


def extract_text_from_docx(file_path):
    """
    Extract all text content from a DOCX file.
    Includes paragraphs and table cell text.
    """
    doc = Document(file_path)
    text_parts = []

    # Extract paragraph text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables (common in resume templates)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts).strip()
